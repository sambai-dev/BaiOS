# SPDX-License-Identifier: AGPL-3.0-or-later
# Attribution and additional terms: see NOTICE.md.

from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "output" / "docx" / "Sam_Bai_Product_Engineer_Resume_2026.docx"

# compact_reference_guide with the named SF_Founding_Engineer_A4 override.
# One font family, one column, explicit spacing, no tables, and no header/footer
# content keep the document visually restrained and ATS-readable.
FONT = "Arial"
INK = RGBColor(0x11, 0x11, 0x11)
MUTED = RGBColor(0x4A, 0x4A, 0x4A)
LINK = RGBColor(0x20, 0x20, 0x20)
RULE = "B9B9B9"

PAGE_WIDTH = Mm(210)
PAGE_HEIGHT = Mm(297)
MARGIN_LEFT_RIGHT = Mm(14)
MARGIN_TOP_BOTTOM = Mm(13)
CONTENT_WIDTH_IN = Mm(210 - 28).inches


def set_run_font(
    run,
    *,
    size: float,
    bold: bool = False,
    italic: bool = False,
    color: RGBColor = INK,
    character_spacing: int | None = None,
):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        r_fonts.set(qn(key), FONT)

    language = OxmlElement("w:lang")
    language.set(qn("w:val"), "en-NZ")
    language.set(qn("w:eastAsia"), "en-NZ")
    r_pr.append(language)

    if character_spacing is not None:
        spacing = OxmlElement("w:spacing")
        spacing.set(qn("w:val"), str(character_spacing))
        r_pr.append(spacing)
    return run


def set_paragraph_spacing(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line: float | None = None,
    exact_line: float | None = None,
):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    if exact_line is not None:
        fmt.line_spacing = Pt(exact_line)
    elif line is not None:
        fmt.line_spacing = line
    fmt.widow_control = True


def set_bottom_border(paragraph, *, color: str = RULE, size: str = "4", space: str = "6"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink(paragraph, text: str, url: str, *, size: float = 9.0):
    relation_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relation_id)

    run = OxmlElement("w:r")
    run_props = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(key), FONT)
    run_props.append(fonts)
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "202020")
    run_props.append(color)
    font_size = OxmlElement("w:sz")
    font_size.set(qn("w:val"), str(int(round(size * 2))))
    run_props.append(font_size)
    font_size_cs = OxmlElement("w:szCs")
    font_size_cs.set(qn("w:val"), str(int(round(size * 2))))
    run_props.append(font_size_cs)
    run.append(run_props)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_separator(paragraph):
    set_run_font(paragraph.add_run("  |  "), size=8.9, color=MUTED)


def add_real_bullet_numbering(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    existing_abstract = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    abstract_id = max(existing_abstract, default=-1) + 1
    existing_nums = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    num_id = max(existing_nums, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)

    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    number_format = OxmlElement("w:numFmt")
    number_format.set(qn("w:val"), "bullet")
    level.append(number_format)
    level_text = OxmlElement("w:lvlText")
    level_text.set(qn("w:val"), "•")
    level.append(level_text)
    level_justification = OxmlElement("w:lvlJc")
    level_justification.set(qn("w:val"), "left")
    level.append(level_justification)

    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "340")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "340")
    indent.set(qn("w:hanging"), "200")
    p_pr.append(indent)
    level.append(p_pr)

    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), FONT)
    fonts.set(qn("w:hAnsi"), FONT)
    r_pr.append(fonts)
    bullet_color = OxmlElement("w:color")
    bullet_color.set(qn("w:val"), "565B61")
    r_pr.append(bullet_color)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)

    number = OxmlElement("w:num")
    number.set(qn("w:numId"), str(num_id))
    abstract_reference = OxmlElement("w:abstractNumId")
    abstract_reference.set(qn("w:val"), str(abstract_id))
    number.append(abstract_reference)
    numbering.append(number)
    return num_id


def apply_bullet(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    level = OxmlElement("w:ilvl")
    level.set(qn("w:val"), "0")
    number = OxmlElement("w:numId")
    number.set(qn("w:val"), str(num_id))
    num_pr.append(level)
    num_pr.append(number)
    p_pr.append(num_pr)


def add_section_heading(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Resume Section")
    set_paragraph_spacing(paragraph, before=9.5, after=3.4, exact_line=11.0)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(
        paragraph.add_run(text.upper()),
        size=9.2,
        bold=True,
        color=INK,
        character_spacing=8,
    )


def add_role_header(doc: Document, role: str, company: str, dates: str):
    paragraph = doc.add_paragraph(style="Resume Role")
    set_paragraph_spacing(paragraph, before=0, after=1.8, exact_line=12.2)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )
    set_run_font(paragraph.add_run(role), size=10.3, bold=True)
    set_run_font(paragraph.add_run(" | "), size=10.3, color=MUTED)
    set_run_font(paragraph.add_run(company), size=10.3, bold=True)
    set_run_font(paragraph.add_run("\t" + dates), size=9.0, bold=True, color=MUTED)


def add_meta_line(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="Resume Meta")
    set_paragraph_spacing(paragraph, before=0, after=3.5, exact_line=10.8)
    paragraph.paragraph_format.keep_with_next = True
    set_run_font(paragraph.add_run(text), size=9.0, color=MUTED)


def add_product_header(doc: Document, name: str, descriptor: str, dates: str):
    paragraph = doc.add_paragraph(style="Resume Product")
    set_paragraph_spacing(paragraph, before=5.0, after=1.8, exact_line=11.5)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )
    set_run_font(paragraph.add_run(name.upper()), size=9.4, bold=True)
    if descriptor:
        set_run_font(paragraph.add_run(" | " + descriptor), size=9.3, color=MUTED)
    set_run_font(paragraph.add_run("\t" + dates), size=8.8, bold=True, color=MUTED)


def add_bullet(doc: Document, text: str, num_id: int):
    paragraph = doc.add_paragraph(style="Resume Bullet")
    set_paragraph_spacing(paragraph, before=0, after=2.3, exact_line=11.5)
    paragraph.paragraph_format.keep_together = True
    apply_bullet(paragraph, num_id)
    set_run_font(paragraph.add_run(text), size=9.35)


def add_skill_line(doc: Document, label: str, text: str):
    paragraph = doc.add_paragraph(style="Resume Detail")
    set_paragraph_spacing(paragraph, before=0, after=1.5, exact_line=11.2)
    paragraph.paragraph_format.keep_together = True
    set_run_font(paragraph.add_run(label + ": "), size=9.2, bold=True)
    set_run_font(paragraph.add_run(text), size=9.2)


def add_education(doc: Document):
    paragraph = doc.add_paragraph(style="Resume Detail")
    set_paragraph_spacing(paragraph, before=0, after=0, exact_line=12.4)
    paragraph.paragraph_format.keep_together = True
    paragraph.paragraph_format.tab_stops.add_tab_stop(
        Inches(CONTENT_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT
    )
    set_run_font(
        paragraph.add_run("Bachelor of Applied Information Technology (Software Engineering)"),
        size=9.4,
        bold=True,
    )
    set_run_font(paragraph.add_run(" | Wintec"), size=9.4, color=MUTED)
    set_run_font(paragraph.add_run("\t2024"), size=8.9, bold=True, color=MUTED)


def configure_styles(doc: Document):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(9.35)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = Pt(11.5)

    existing = {style.name for style in doc.styles}
    for name in (
        "Resume Section",
        "Resume Role",
        "Resume Meta",
        "Resume Product",
        "Resume Bullet",
        "Resume Detail",
    ):
        if name not in existing:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style = doc.styles[name]
        style.base_style = normal
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(9.35)


def build():
    doc = Document()
    configure_styles(doc)

    section = doc.sections[0]
    section.page_width = PAGE_WIDTH
    section.page_height = PAGE_HEIGHT
    section.top_margin = MARGIN_TOP_BOTTOM
    section.bottom_margin = MARGIN_TOP_BOTTOM
    section.left_margin = MARGIN_LEFT_RIGHT
    section.right_margin = MARGIN_LEFT_RIGHT
    section.header_distance = Mm(5)
    section.footer_distance = Mm(5)

    properties = doc.core_properties
    properties.title = "Sam Bai - Product Engineer Resume"
    properties.subject = "One-page startup-focused product engineering resume"
    properties.author = "Sam Bai"
    properties.keywords = (
        "Product Engineer, Full Stack Engineer, Founding Engineer, TypeScript, React, "
        "Next.js, PostgreSQL, AI systems, New Zealand"
    )

    num_id = add_real_bullet_numbering(doc)

    # memo_masthead adapted for an ATS-safe resume: identity, positioning and
    # contact details remain in the body, with a single restrained divider.
    name = doc.add_paragraph()
    set_paragraph_spacing(name, before=0, after=0.2, exact_line=23.0)
    set_run_font(name.add_run("SAM BAI"), size=21.5, bold=True)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, before=0, after=2.0, exact_line=13.0)
    set_run_font(
        title.add_run("PRODUCT ENGINEER"),
        size=11.5,
        bold=True,
        color=INK,
        character_spacing=2,
    )

    contact = doc.add_paragraph()
    set_paragraph_spacing(contact, before=0, after=0.5, exact_line=11.0)
    set_run_font(contact.add_run("Hamilton, New Zealand"), size=8.5)
    add_separator(contact)
    set_run_font(contact.add_run("NZ Permanent Resident"), size=8.5)
    add_separator(contact)
    set_run_font(contact.add_run("+64 27 460 4700"), size=8.5)
    add_separator(contact)
    add_hyperlink(contact, "sambai.codes@gmail.com", "mailto:sambai.codes@gmail.com")

    links = doc.add_paragraph()
    set_paragraph_spacing(links, before=0, after=6.5, exact_line=10.8)
    add_hyperlink(links, "sambai.dev", "https://www.sambai.dev")
    add_separator(links)
    add_hyperlink(links, "github.com/sambai-dev", "https://github.com/sambai-dev")
    add_separator(links)
    add_hyperlink(
        links,
        "linkedin.com/in/sam-bai1/",
        "https://www.linkedin.com/in/sam-bai1/",
    )
    set_bottom_border(links)

    add_section_heading(doc, "Experience")
    add_role_header(
        doc,
        "Founder & Product Engineer",
        "Solynth Labs Limited",
        "2026 - Present",
    )
    add_meta_line(
        doc,
        "Founded and run a New Zealand software studio; lead product discovery, architecture, implementation, deployment and client support.",
    )

    add_product_header(doc, "Rivet", "Construction operations - client project", "Jul 2026 - Present")
    add_bullet(
        doc,
        "Built and deployed a private construction-operations PWA for a New Zealand residential builder, joining project, task, exception, approval, calendar and notification workflows for the director and project managers.",
        num_id,
    )
    add_bullet(
        doc,
        "Designed server-authoritative schedule proposals that recalculate before apply, reject stale changes and write dates, audit records and notification effects atomically in PostgreSQL.",
        num_id,
    )
    add_bullet(
        doc,
        "Implemented scoped authentication and conflict-aware browser sync with IndexedDB outboxes, version checks and transactional Hono APIs on Cloudflare Workers.",
        num_id,
    )

    add_product_header(doc, "Trekky", "Job-search product", "Jan 2026 - Present")
    add_bullet(
        doc,
        "Built and operate a live Next.js/TypeScript product for job discovery, application tracking, contacts, AI preparation, outcome analytics and review-first automation.",
        num_id,
    )
    add_bullet(
        doc,
        "Added account-scoped MCP access, a Chrome MV3 capture extension and multi-source ingestion across NZ, Australia, the US and Singapore with duplicate detection and source history.",
        num_id,
    )
    add_bullet(
        doc,
        "Cut default dashboard JavaScript from 341 KB to 233 KB gzip through route-level lazy loading.",
        num_id,
    )

    add_section_heading(doc, "Selected Projects")
    add_product_header(doc, "Clearfold", "Pre-production digital-asset operations", "2026")
    add_bullet(
        doc,
        "Built exact-decimal TypeScript libraries for double-entry ledger rules, FIFO lots, portfolio valuation and versioned risk evaluation, with idempotency and reversal invariants.",
        num_id,
    )
    add_bullet(
        doc,
        "Implemented a private Gemini assistant with cited conversation history and pgvector retrieval, with lexical fallback when embeddings are unavailable.",
        num_id,
    )

    add_product_header(doc, "BaiOS Workbench", "Open-source browser interface", "Dec 2025 - Present")
    add_bullet(
        doc,
        "Built and deployed sambai.dev as a 16-app, three-workspace Workbench with multi-instance windows, drag, resize, snap, Atlas navigation and a searchable local archive.",
        num_id,
    )
    add_bullet(
        doc,
        "Designed atomically committed browser-local state, cross-tab revision-conflict handling and schema-validated backup/import without implying remote sync.",
        num_id,
    )

    add_section_heading(doc, "Open Source Contributions")
    add_product_header(doc, "EdgarTools & Betaflight Configurator", "", "2026")
    add_bullet(
        doc,
        "Built Blackbox MP4/WebM video export for Betaflight Configurator using WebCodecs, with overlay compositing, progress/ETA, cancellation and streaming saves; merged as PR #5468 for 2026.12.",
        num_id,
    )
    add_bullet(
        doc,
        "Resolved a load-dependent Vitest teardown race through upstream CI and maintainer review; merged for Betaflight Configurator 2026.12.",
        num_id,
    )
    add_bullet(
        doc,
        "Repaired legacy SEC 13F TXT parsing so checksum-valid CUSIPs survive column drift; accepted upstream and released in EdgarTools v5.53.0.",
        num_id,
    )

    add_section_heading(doc, "Technical Skills")
    add_skill_line(
        doc,
        "Languages & frameworks",
        "TypeScript, JavaScript, Python, SQL, React, Next.js, Node.js, Vite, Hono, Tailwind CSS, TanStack Query, Router, Table and Virtual",
    )
    add_skill_line(
        doc,
        "Data & systems",
        "PostgreSQL (Neon), MongoDB Atlas, Mongoose, Redis, Drizzle ORM, IndexedDB/Dexie, role/capability authorisation, workers/queues, MCP, vector/lexical retrieval",
    )
    add_skill_line(
        doc,
        "Cloud & testing",
        "Cloudflare Workers, R2, Hyperdrive, Vercel, GitHub Actions, Vitest, Playwright",
    )

    add_section_heading(doc, "Education")
    add_education(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    try:
        build()
    except Exception as error:  # Surface a clean failure instead of a raw traceback.
        print(f"resume build failed: {error}", file=sys.stderr)
        raise SystemExit(1) from error
